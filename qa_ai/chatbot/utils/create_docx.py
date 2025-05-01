from pathlib import Path
from app.models.question import QAResponse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import HTTPException
import re


def create_formatted_docx_file(
    qa_result: QAResponse, docx_file_path: Path, exam_subject: str, exam_duration: str
) -> None:
    doc = Document()

    # Thiết lập font chữ mặc định
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Phần tiêu đề
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run("ĐỀ THI TỰ LUẬN")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Thông tin môn thi & thời gian
    info_paragraph = doc.add_paragraph()
    info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_paragraph.add_run(
        f"Môn thi: {exam_subject}\nThời gian làm bài: {exam_duration}"
    )
    info_run.bold = True
    info_run.font.size = Pt(12)

    # Hướng dẫn
    guide_paragraph = doc.add_paragraph()
    guide_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    guide_run = guide_paragraph.add_run("(Thí sinh không được phép sử dụng tài liệu)")
    guide_run.italic = True
    guide_run.font.size = Pt(11)

    doc.add_paragraph("")

    # Phân bố cấp độ Bloom
    doc.add_heading("PHÂN BỐ CẤP ĐỘ BLOOM:", level=1)
    if qa_result.bloom_assignment:
        doc.add_paragraph(qa_result.bloom_assignment)
    else:
        doc.add_paragraph("Không có phân bổ cấp độ Bloom.")

    num_levels = len(qa_result.qa_results)
    percentages = []

    if num_levels == 0:
        percentages = []
    elif num_levels == 1:
        percentages = [1.0]
    elif num_levels == 6:
        percentages = [0.10, 0.15, 0.20, 0.20, 0.20, 0.15]
    else:
        percentages = [0.10] + [0.90 / (num_levels - 1)] * (num_levels - 1)

    # Câu hỏi và câu trả lời
    doc.add_heading("CÂU HỎI VÀ CÂU TRẢ LỜI:", level=1)
    question_counter = 1

    for level_idx, level_result in enumerate(qa_result.qa_results):
        doc.add_heading(level_result["level"], level=2)
        questions_dict = level_result["questions"]
        num_questions = len(questions_dict)

        level_total_points = 0
        if level_idx < len(percentages):
            level_total_points = percentages[level_idx] * 10

        points_alloc_per_question = []
        if num_questions > 0 and level_total_points > 0:
            for i in range(num_questions - 1):
                raw = level_total_points / num_questions
                nice = round(raw * 20) / 20
                points_alloc_per_question.append(nice)
            sum_so_far = sum(points_alloc_per_question)
            last_point = level_total_points - sum_so_far
            last_nice = round(last_point * 20) / 20
            points_alloc_per_question.append(last_nice)
        else:
            points_alloc_per_question = [0.0] * num_questions

        for idx, (question, answer) in enumerate(questions_dict.items()):
            clean_question = re.sub(r"^Câu \d+:\s*", "", question.strip())
            formatted_question = f"Câu {question_counter}: {clean_question}"

            para_question = doc.add_paragraph()
            run_question_text = para_question.add_run(
                clean_markdown_bold(formatted_question)
            )
            run_question_text.bold = True

            points_display = 0.0
            if idx < len(points_alloc_per_question):
                points_display = points_alloc_per_question[idx]

            run_points = para_question.add_run(f" ({points_display:.2f} điểm)")
            run_points.bold = True

            para_answer_title = doc.add_paragraph()
            run_answer_title = para_answer_title.add_run("Trả lời: ")
            run_answer_title.bold = True

            para_answer_content = doc.add_paragraph(clean_markdown_bold(answer))
            para_answer_content.left_indent = Inches(0.5)

            question_counter += 1

    doc.save(docx_file_path)


def create_simple_docx_file(
    qa_result: QAResponse, docx_file_path: Path, exam_subject: str, exam_duration: str
) -> None:
    """
    Tạo file DOCX đơn giản chỉ chứa phần header và danh sách câu hỏi với điểm,
    sử dụng cơ chế phân bố điểm giống như create_formatted_docx_file.
    """
    doc = Document()

    # ------------------- PHẦN HEADER -------------------
    # (Giữ nguyên phần header)
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run("ĐỀ THI TỰ LUẬN")
    title_run.bold = True
    title_run.font.size = Pt(14)

    info_paragraph = doc.add_paragraph()
    info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_paragraph.add_run(
        f"Môn thi: {exam_subject}\nThời gian làm bài: {exam_duration}"
    )
    info_run.bold = True
    info_run.font.size = Pt(12)

    guide_paragraph = doc.add_paragraph()
    guide_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    guide_run = guide_paragraph.add_run("(Thí sinh không được phép sử dụng tài liệu)")
    guide_run.italic = True
    guide_run.font.size = Pt(11)

    doc.add_paragraph("")

    # Xác định số cấp độ và TÍNH TOÁN TỶ LỆ ĐIỂM MỚI
    num_levels = len(qa_result.qa_results)
    percentages = []  # Khởi tạo danh sách tỷ lệ điểm

    if num_levels == 0:
        # Không có cấp độ nào có câu hỏi, tỷ lệ điểm rỗng
        percentages = []
    elif num_levels == 1:
        # Chỉ có một cấp độ, cấp độ này nhận 100% điểm (10 điểm)
        percentages = [1.0]
    elif num_levels == 6:
        # Tỷ lệ điểm cố định cho trường hợp có chính xác 6 cấp độ
        percentages = [0.10, 0.15, 0.20, 0.20, 0.20, 0.15]
    else:  # num_levels > 1 và num_levels != 6 (tức là 2, 3, 4, 5 cấp độ)
        # 10% cho cấp độ đầu tiên, 90% chia đều cho các cấp độ còn lại
        percentages = [0.10] + [0.90 / (num_levels - 1)] * (num_levels - 1)

    # Tính toán và phân phối điểm cho TỪNG CÂU HỎI
    question_counter = 1  # Bộ đếm câu hỏi tổng

    for level_idx, level_result in enumerate(qa_result.qa_results):
        questions_dict = level_result["questions"]
        num_questions = len(questions_dict)

        # Tính tổng điểm cho cấp dựa trên tỷ lệ đã tính
        level_total_points = 0
        if level_idx < len(percentages):
            level_total_points = percentages[level_idx] * 10

        points_alloc_per_question = []
        if (
            num_questions > 0 and level_total_points > 0
        ):  # Chỉ phân bổ nếu có câu hỏi và tổng điểm > 0
            for i in range(num_questions - 1):
                raw = level_total_points / num_questions
                nice = round(raw * 20) / 20
                points_alloc_per_question.append(nice)
            sum_so_far = sum(points_alloc_per_question)
            last_point = level_total_points - sum_so_far
            last_nice = round(last_point * 20) / 20
            points_alloc_per_question.append(last_nice)
        else:
            points_alloc_per_question = [
                0.0
            ] * num_questions  # Gán 0 điểm cho mọi câu nếu không có điểm

        # Duyệt qua từng câu hỏi trong cấp độ và thêm vào doc
        for idx, (question, _) in enumerate(
            questions_dict.items()
        ):  # simple file không cần answer
            # Loại bỏ tiền tố "Câu X:" nếu đã tồn tại (giữ lại logic này)
            clean_question = re.sub(r"^Câu \d+:\s*", "", question.strip())
            # Định dạng câu hỏi với tiền tố đúng (đánh số liên tục)
            formatted_question = f"Câu {question_counter}: {clean_question}"

            # Thêm câu hỏi và điểm vào đoạn văn mới
            para = doc.add_paragraph()
            points_display = 0.0  # Điểm mặc định là 0
            if idx < len(points_alloc_per_question):
                points_display = points_alloc_per_question[idx]

            para.add_run(
                f"{formatted_question} ({points_display:.2f} điểm)"
            )  # Định dạng 2 chữ số sau thập phân

            # Tăng bộ đếm câu hỏi tổng
            question_counter += 1

    # Lưu tài liệu
    doc.save(docx_file_path)


def clean_markdown_bold(text: str) -> str:
    return re.sub(r"\*\*(.*?)\*\*", r"\1", text)
