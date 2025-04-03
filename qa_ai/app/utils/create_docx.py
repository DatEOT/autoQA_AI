from pathlib import Path
from app.models.question import QAResponse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_formatted_docx_file(
    qa_result: QAResponse, docx_file_path: Path, exam_subject: str, exam_duration: str
) -> None:
    """
    Tạo file DOCX với định dạng:
      - Tiêu đề: "ĐỀ THI TỰ LUẬN" (gán cứng)
      - Môn thi và thời gian làm bài: do người dùng nhập vào
      - Phân bổ điểm cho mỗi cấp độ Bloom (từ thấp đến cao) và chia đều số điểm cho từng câu hỏi trong mỗi cấp,
        sao cho tổng số điểm là 10.
        Ví dụ: Nếu cấp được 3 điểm và có 2 câu hỏi, mỗi câu nhận 1.5 điểm; nếu có 7 câu thì mỗi câu nhận ~0.43 điểm.
      - Danh sách câu hỏi (theo cấp độ Bloom) kèm đáp án
    """
    doc = Document()

    # -- Thiết lập font chữ mặc định --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # ------------------- PHẦN TIÊU ĐỀ -------------------
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run("ĐỀ THI TỰ LUẬN")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ------------------- THÔNG TIN MÔN THI & THỜI GIAN -------------------
    info_paragraph = doc.add_paragraph()
    info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_paragraph.add_run(
        f"Môn thi: {exam_subject}\nThời gian làm bài: {exam_duration}"
    )
    info_run.bold = True
    info_run.font.size = Pt(12)

    # ------------------- HƯỚNG DẪN -------------------
    guide_paragraph = doc.add_paragraph()
    guide_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    guide_run = guide_paragraph.add_run("(Thí sinh không được phép sử dụng tài liệu)")
    guide_run.italic = True
    guide_run.font.size = Pt(11)

    # Dòng trống ngăn cách
    doc.add_paragraph("")

    # ------------------- PHÂN BỔ CẤP ĐỘ BLOOM -------------------
    doc.add_heading("PHÂN BỔ CẤP ĐỘ BLOOM:", level=1)
    doc.add_paragraph(qa_result.bloom_assignment)

    # Tính số cấp độ và tổng trọng số (cấp 1 có trọng số 1, cấp 2 trọng số 2, v.v.)
    num_levels = len(qa_result.qa_results)
    total_weight = sum(range(1, num_levels + 1))

    # Tính điểm cho từng cấp dưới dạng số thực với tổng điểm là 10
    # raw_points = 10 * (i+1) / total_weight cho cấp i
    level_points_alloc = [10 * (i + 1) / total_weight for i in range(num_levels)]

    # ------------------- CÂU HỎI VÀ CÂU TRẢ LỜI -------------------
    doc.add_heading("CÂU HỎI VÀ CÂU TRẢ LỜI:", level=1)
    question_counter = 1

    # Lặp qua các cấp độ Bloom (giả sử thứ tự trong danh sách là từ thấp đến cao)
    for level_idx, level_result in enumerate(qa_result.qa_results):
        # Heading cho cấp độ Bloom (ví dụ: "Cấp độ 1 - Nhớ")
        doc.add_heading(level_result["level"], level=2)
        questions_dict = level_result["questions"]
        num_questions = len(questions_dict)

        # Tính điểm được phân bổ cho cấp hiện tại
        level_total_points = level_points_alloc[level_idx]
        # Chia đều số điểm cho mỗi câu (làm tròn đến 2 chữ số thập phân)
        points_per_question = (
            round(level_total_points / num_questions, 2) if num_questions > 0 else 0
        )

        for question, answer in questions_dict.items():
            para_question = doc.add_paragraph()
            run_question_title = para_question.add_run(f"Câu {question_counter}: ")
            run_question_title.bold = True

            run_question_text = para_question.add_run(question)
            run_question_text.font.bold = False

            run_points = para_question.add_run(f" ({points_per_question} điểm)")
            run_points.bold = True

            para_answer_title = doc.add_paragraph()
            run_answer_title = para_answer_title.add_run("Trả lời: ")
            run_answer_title.bold = True

            para_answer_content = doc.add_paragraph(answer)
            para_answer_content.left_indent = Inches(0.5)

            question_counter += 1

    doc.save(docx_file_path)


def create_simple_docx_file(
    qa_result: QAResponse, docx_file_path: Path, exam_subject: str, exam_duration: str
) -> None:
    """
    Tạo file DOCX đơn giản chỉ chứa phần header giống file formatted_docx_file và danh sách câu hỏi với điểm.
    Header bao gồm:
      - ĐỀ THI TỰ LUẬN
      - Môn thi: {exam_subject}
      - Thời gian làm bài: {exam_duration}
      - (Thí sinh không được phép sử dụng tài liệu)
    Sau đó là danh sách các câu hỏi với điểm được chia đều theo cấp độ Bloom (từ thấp đến cao) sao cho tổng điểm là 10.
    Mỗi câu hỏi nhận số điểm là:
          level_total_points / số câu hỏi trong cấp
    (Ví dụ: nếu cấp được 3 điểm và có 2 câu hỏi, mỗi câu nhận 1.5 điểm.)
    """
    doc = Document()

    # ------------------- PHẦN HEADER -------------------
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run("ĐỀ THI TỰ LUẬN")
    title_run.bold = True
    title_run.font.size = Pt(14)

    info_paragraph = doc.add_paragraph()
    info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_paragraph.add_run(
        f"Môn thi: {exam_subject}\nThời gian làm bài: {exam_duration}"
    )
    info_run.bold = True
    info_run.font.size = Pt(12)

    guide_paragraph = doc.add_paragraph()
    guide_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    guide_run = guide_paragraph.add_run("(Thí sinh không được phép sử dụng tài liệu)")
    guide_run.italic = True
    guide_run.font.size = Pt(11)

    doc.add_paragraph("")

    # Tính số cấp độ và tổng trọng số
    num_levels = len(qa_result.qa_results)
    total_weight = sum(range(1, num_levels + 1))

    # Tính điểm cho từng cấp dưới dạng số thực với tổng điểm là 10
    level_points_alloc = [10 * (i + 1) / total_weight for i in range(num_levels)]

    question_counter = 1

    for level_idx, level_result in enumerate(qa_result.qa_results):
        questions_dict = level_result["questions"]
        num_questions = len(questions_dict)
        level_total_points = level_points_alloc[level_idx]
        points_per_question = (
            round(level_total_points / num_questions, 2) if num_questions > 0 else 0
        )

        for question, _ in questions_dict.items():
            para = doc.add_paragraph()
            para.add_run(
                f"Câu {question_counter}: {question} ({points_per_question} điểm)"
            )
            question_counter += 1

    doc.save(docx_file_path)
