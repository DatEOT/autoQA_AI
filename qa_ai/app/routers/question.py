import os
import logging
from fastapi import APIRouter, UploadFile, File, HTTPException, Form, Depends
from app.models.question import QuestionRequest, FileResponseModel, QAResponse
from docx import Document
import PyPDF2
from uuid import uuid4
import openai
from dotenv import load_dotenv
import httpx
from app.security.security import get_api_key
from app.utils.ingestion import Ingestion
from typing import List, Tuple, Dict
from pathlib import Path
import re
from app.utils.create_docx import create_formatted_docx_file
from app.utils.create_docx import create_simple_docx_file
from app.utils.create_pdf import create_formatted_pdf_file
from app.utils.create_pdf import create_simple_pdf_file
from fastapi.responses import StreamingResponse
import io
import zipfile
from app.utils.mysql_connection import get_db
import pymysql
from app.security.dependency import get_current_user_id
from docx2pdf import convert

# Cấu hình logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Tải biến môi trường
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    raise ValueError("API key không được tìm thấy trong biến môi trường")

# Khởi tạo OpenAI client với timeout
client = openai.OpenAI(
    api_key=OPENAI_API_KEY,
    http_client=httpx.Client(timeout=httpx.Timeout(30.0)),
)

router = APIRouter(prefix="/questions", tags=["questions"])

# Cấu hình thư mục
UPLOAD_DIR = Path("qa_ai/uploads")
TXT_DIR = Path("qa_ai/txt")

# Đảm bảo thư mục tồn tại
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
TXT_DIR.mkdir(parents=True, exist_ok=True)

# Bloom taxonomy keywords
BLOOM_KEYWORDS = {
    1: "dán nhãn, định nghĩa, đọc, gọi tên, lặp lại, nghe hiểu, nhắc lại, nhận biết, nhận dạng, nhận diện, nhận ra, tìm ra, nối kết điểm phù hợp, nêu rõ, chép lại, sắp xếp, sắp xếp theo thứ tự, kể lại, tóm tắt lại, trích dẫn, nêu ra, trình bày đại ý, ghi nhớ, mô tả, liệt kê, kể tên, ghi lại, đọc lại, thuật lại, tìm câu trả lời, kể lại, nhắc lại, xác định, tìm, chọn, sao chép, lặp lại chính xác, lập bảng, mô phỏng, sao chép, kể, làm theo, tưởng tượng, hình dung, loại bỏ, xem xét, kiểm tra chi tiết, tìm câu trả lời",
    2: "giải thích, diễn giải, tóm tắt, so sánh, phân biệt, hỏi, đặt vấn đề, liên hệ, trích dẫn, tìm xuất xứ, phân loại, tính toán, ước tính, so sánh, đối chiếu, chuyển đổi, xác định, nhận dạng, chỉ ra, suy ra, phỏng đoán, đánh giá, nhận định, xếp thứ tự, mô tả, khám phá, tìm ra, phát hiện, trình bày & phân tích, thảo luận, nhận biết, ước lượng, đánh giá, phán đoán, dự đoán, thể hiện, trình bày, diễn đạt, mở rộng, khái quát hóa, tổng hợp, đưa ra, nêu ví dụ, phân nhóm, sắp xếp theo nhóm, nhận diện, báo cáo, nghiên cứu, tìm, trình bày lại, nhắc lại, xem lại, viết lại, lựa chọn, tóm lược, phát hiện, chỉ ra (nguồn gốc), tìm xuất xứ, diễn dịch, dịch ra, chuyển sang, chuyển ngữ",
    3: "áp dụng, sử dụng, thực hiện, minh họa, tính toán, vận dụng, điều khiển, xử lý, thêm vào, sửa đổi, điều chỉnh, thay đổi, vận hành, chọn lựa, chọn lọc, lựa chọn, thực hành, phân loại, phân loại chủ đề, ước lượng, dự đoán, dự báo, hoàn thành, chuẩn bị, ước tính, tạo ra, xây dựng, liên hệ kiến thức, thuyết minh, chứng minh, giải thích, biểu diễn, lập kế hoạch, soạn/viết thành kịch bản, biên kịch, chỉ ra, nêu ra, phác họa, giải quyết, khái quát hóa, tổng quát hóa, loại trừ, sơ đồ hóa, thể hiện đồ thị, vẽ biểu đồ, thêm thông tin, chia nhỏ, diễn giải, làm sáng tỏ, dịch, liệt kê, đánh giá, nhận định, mô tả (âm thanh...)",
    4: "phân tích, phân loại, đánh giá, xác định, so sánh chi tiết, sắp xếp, chia nhỏ, phân nhỏ, tách nhỏ ra, tính toán, phạm trù hóa, phân nhóm, thay đổi, lựa chọn, kết hợp, so sánh, ước tính, so sánh đối chiếu, phản biện, phê bình, minh họa, bình luận, chứng minh, giải thích, thiết kế, phân biệt, khu biệt, khám phá, tìm ra, phát hiện, soạn/viết thành kịch bản, xem xét, kiểm tra, thử nghiệm, suy luận, điều hành, vận dụng, xử lý, mô hình hóa, điều chỉnh, sửa đổi, vận hành, điều hành, lập dàn ý, nêu đại ý, phác thảo, phát hiện, nhận ra, chỉ ra, phát triển, thực hiện, thực hành, thể hiện đồ thị, sơ đồ hóa, ước đoán, dự đoán, dự báo, đặt câu hỏi, đặt vấn đề, tạo ra, chế tạo, sắp xếp kế hoạch, liên hệ, tách ra, khảo sát, phân chia, kiểm tra, thử nghiệm",
    5: "đánh giá, phê bình, biện minh, bảo vệ, đưa ra ý kiến, thẩm định, sắp xếp, thu thập, tập hợp lại, phối hợp, phân loại, phân chia, kết nối, so sánh, kết hợp, sáng tác, kết luận, xây dựng, tạo ra, tranh luận, xét đoán, phát triển, thiết kế, xác lập công thức, suy ra, đặt vấn đề, đặt giả thiết, nhận định, điều chỉnh, chuẩn bị, sắp xếp lại, xây dựng lại, lặp lại, liên hệ (tổng hợp), tái cấu trúc, ra soát, xem xét lại, sửa lại, viết lại, thiết lập, lập (kế hoạch, mô hình…), tổng kết, tóm tắt, cho ý kiến, định giá, cân nhắc, xem xét cẩn thận, viết (báo cáo, dự thảo…), tổ chức, sắp xếp theo hệ thống",
    6: "thiết kế, sáng tạo, xây dựng, phát triển, đề xuất, tranh luận, biện luận, kết hợp, sáng tác, lập, thiết lập, tích hợp, diễn dịch, diễn giải, tìm ra, phát minh ra, làm ra, đối chiếu, lập kế hoạch, tạo ra, chế tạo, đưa ra, đặt ra, phê bình, bình luận, tổng hợp, ước lượng, biến đổi, sáng chế, đề nghị",
}


def save_uploaded_file(file: UploadFile, file_id: str) -> Path:
    """Lưu file tải lên và trả về đường dẫn."""
    try:
        file_path = UPLOAD_DIR / f"{file_id}_{file.filename}"
        with open(file_path, "wb") as buffer:
            buffer.write(file.file.read())
        return file_path
    except Exception as e:
        logger.error(f"Lỗi khi lưu file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Lỗi khi lưu file: {str(e)}")


def extract_text_from_file(file_path: Path, file_extension: str) -> str:
    """Trích xuất văn bản từ file dựa trên loại file."""
    try:
        if file_extension == "docx":
            doc = Document(file_path)
            return "\n".join(
                [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            )
        elif file_extension == "pdf":
            with open(file_path, "rb") as pdf_file:
                reader = PyPDF2.PdfReader(pdf_file)
                texts = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        texts.append(text.strip())
                return "\n".join(texts)
        elif file_extension == "txt":
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        raise ValueError("Loại file không được hỗ trợ")
    except Exception as e:
        logger.error(f"Lỗi khi trích xuất văn bản: {str(e)}")
        raise HTTPException(
            status_code=500, detail=f"Lỗi khi trích xuất văn bản: {str(e)}"
        )


def validate_request(request: QuestionRequest) -> None:
    """Kiểm tra tính hợp lệ của yêu cầu."""
    total_levels = sum(
        [
            request.level_1,
            request.level_2,
            request.level_3,
            request.level_4,
            request.level_5,
            request.level_6,
        ]
    )

    if total_levels == 0:
        raise HTTPException(
            status_code=400,
            detail="Không có câu hỏi nào được yêu cầu ở bất kỳ cấp độ nào.",
        )

    if total_levels != request.num_questions:
        raise HTTPException(
            status_code=400,
            detail=f"Tổng số câu hỏi theo cấp độ ({total_levels}) không khớp với num_questions ({request.num_questions})",
        )


def get_bloom_level_name(level: int) -> str:
    """Lấy tên cấp độ Bloom tương ứng."""
    level_names = {
        1: "Nhớ",
        2: "Hiểu",
        3: "Áp dụng",
        4: "Phân tích",
        5: "Đánh giá",
        6: "Sáng tạo",
    }
    return level_names.get(level, "Không xác định")


def get_bloom_level_description(level: int) -> str:
    """Lấy mô tả cho từng cấp độ Bloom."""
    descriptions = {
        1: "Nhớ lại thông tin cơ bản, dữ liệu.",
        2: "Giải thích hoặc làm rõ khái niệm.",
        3: "Áp dụng kiến thức vào thực tế.",
        4: "Phân tích và giải thích mối quan hệ giữa các yếu tố.",
        5: "Đánh giá và phê phán ý tưởng hoặc giả thuyết.",
        6: "Sáng tạo hoặc thiết kế giải pháp mới.",
    }
    return descriptions.get(level, "Không xác định")


def generate_bloom_assignment(
    segments: List[Tuple[str, str]],
) -> Tuple[List[str], List[int]]:
    """Tạo phân bổ cấp độ Bloom cho các đoạn văn."""
    total_segments = len(segments)
    base_segments_per_level = total_segments // 6
    remainder = total_segments % 6

    segments_per_level = [base_segments_per_level] * 6
    for i in range(remainder):
        segments_per_level[i] += 1

    bloom_assignment = []
    segment_idx = 0
    for level in range(1, 7):
        num_segments = segments_per_level[level - 1]
        level_name = get_bloom_level_name(level)
        for _ in range(num_segments):
            bloom_assignment.append(
                f"- Đoạn văn {segment_idx + 1}: Cấp độ {level} - {level_name}"
            )
            segment_idx += 1

    return bloom_assignment, segments_per_level


def extract_questions_from_response(response: str) -> List[str]:
    """Trích xuất danh sách câu hỏi từ response của OpenAI."""
    questions = []
    lines = response.split("\n")
    for line in lines:
        if re.match(
            r"^\d+\.", line.strip()
        ):  # Match lines starting with "1.", "2.", etc.
            question = re.sub(r"^\d+\.\s*", "", line.strip())  # Remove numbering
            questions.append(question)
    return questions


def generate_answers_for_questions(
    questions: List[str], context: str
) -> Dict[str, str]:
    """Tạo câu trả lời cho các câu hỏi dựa trên nội dung file input."""
    qa_pairs = {}

    for question in questions:
        prompt = (
            f"Dựa trên nội dung sau:\n\n{context}\n\n"
            f"Hãy trả lời câu hỏi: {question}\n"
            "Yêu cầu:\n"
            "1. Câu trả lời phải nằm trong nội dung đã cung cấp\n"
            "2. Nếu không tìm thấy thông tin, trả lời 'Không có thông tin trong tài liệu'\n"
            "3. Trả lời ngắn gọn, súc tích\n\n"
            "Câu trả lời:"
        )

        try:
            response = (
                client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.3,  # Giảm temperature để câu trả lời chính xác hơn
                    max_tokens=500,
                )
                .choices[0]
                .message.content.strip()
            )

            qa_pairs[question] = response
        except Exception as e:
            logger.error(f"Lỗi khi tạo câu trả lời: {str(e)}")
            qa_pairs[question] = "Không thể tạo câu trả lời"

    return qa_pairs


def generate_questions_for_level(
    level: int, segments: List[Tuple[str, str]], num_questions: int, context: str
) -> Dict[str, Dict[str, str]]:
    """Tạo câu hỏi và câu trả lời cho một cấp độ Bloom."""
    if num_questions <= 0:
        return {}

    level_name = get_bloom_level_name(level)
    bloom_keywords = BLOOM_KEYWORDS.get(level, "")
    level_description = get_bloom_level_description(level)

    # Lấy các đoạn văn cho cấp độ này
    level_segments = segments[:num_questions]

    # Tạo danh sách đoạn văn
    segment_texts = []
    for i, (_, text) in enumerate(level_segments):
        segment_texts.append(f'- Đoạn văn {i + 1}: "{text}"')
    level_segment_texts = "\n".join(segment_texts)

    # Tạo prompt cho câu hỏi
    prompt_parts = [
        "Dựa trên các đoạn văn bản sau:",
        level_segment_texts,
        "",
        "Hãy thực hiện bước sau:",
        f"1. Tạo chính xác {num_questions} câu hỏi tự luận thuộc Cấp độ {level} - {level_name}:",
        f"   - Sử dụng từ khóa: {bloom_keywords}",
        "   - Đảm bảo câu hỏi phù hợp với yêu cầu của cấp độ:",
        f"     - {level_description}",
        "",
        "Yêu cầu:",
        "   - Mỗi câu hỏi phải bắt đầu bằng số thứ tự (1., 2., ...)",
        "   - Chỉ trả về danh sách câu hỏi, không thêm giải thích",
    ]

    prompt = "\n".join(prompt_parts)

    try:
        # Tạo câu hỏi
        response = (
            client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7,
            )
            .choices[0]
            .message.content.strip()
        )

        # Trích xuất câu hỏi từ response
        questions = extract_questions_from_response(response)

        # Tạo câu trả lời cho các câu hỏi
        qa_pairs = generate_answers_for_questions(questions, context)

        return {"level": f"Cấp độ {level} - {level_name}", "questions": qa_pairs}
    except Exception as e:
        logger.error(f"Lỗi khi gọi OpenAI API: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Lỗi khi tạo câu hỏi: {str(e)}")


def generate_qa_content(
    segments: List[Tuple[str, str]], request: QuestionRequest, context: str
) -> QAResponse:
    """Tạo nội dung câu hỏi và câu trả lời."""
    # Tạo phân bổ cấp độ Bloom
    bloom_assignment, _ = generate_bloom_assignment(segments)
    bloom_assignment_text = "\n".join(bloom_assignment)

    # Tạo câu hỏi và câu trả lời cho từng cấp độ
    qa_results = []
    for level in range(1, 7):
        num_questions = getattr(request, f"level_{level}")
        if num_questions <= 0:
            continue

        level_result = generate_questions_for_level(
            level=level, segments=segments, num_questions=num_questions, context=context
        )
        if level_result:
            qa_results.append(level_result)

    # Tổng hợp kết quả
    return QAResponse(bloom_assignment=bloom_assignment_text, qa_results=qa_results)


def convert_txt_to_docx_and_delete(txt_path: Path) -> Path:
    """
    Chuyển đổi file txt sang file docx và xóa file txt cũ.

    Args:
        txt_path (Path): Đường dẫn tới file txt cần chuyển đổi.

    Returns:
        Path: Đường dẫn tới file docx mới được tạo.
    """
    # Đọc nội dung từ file txt
    with open(txt_path, "r", encoding="utf-8") as f:
        content = f.read()

    # Tạo file docx mới với cùng tên nhưng đuôi .docx
    docx_file_path = txt_path.with_suffix(".docx")
    doc = Document()

    # Nếu muốn định dạng lại nội dung theo mục lục hoặc chia đoạn, có thể thực hiện thêm xử lý ở đây.
    doc.add_paragraph(content)
    doc.save(docx_file_path)

    # Xóa file txt cũ sau khi chuyển đổi
    os.remove(txt_path)

    return docx_file_path


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    """
    Chuyển đổi file DOCX sang PDF.

    Args:
        docx_path (Path): Đường dẫn tới file DOCX.
        pdf_path (Path): Đường dẫn file PDF sẽ được tạo ra.
    """
    try:
        convert(str(docx_path), str(pdf_path))
    except Exception as e:
        logger.error(f"Lỗi khi chuyển đổi {docx_path} sang PDF: {str(e)}")
        raise HTTPException(
            status_code=500, detail=f"Lỗi khi chuyển đổi file: {str(e)}"
        )


@router.get("/download/zip/{file_id}", response_class=StreamingResponse)
def download_zip(file_id: str):
    # Xác định đường dẫn 4 file: 2 file DOCX và 2 file PDF
    formatted_docx_file = TXT_DIR / f"{file_id}_formatted.docx"
    simple_docx_file = TXT_DIR / f"{file_id}_simple.docx"
    formatted_pdf_file = TXT_DIR / f"{file_id}_formatted.pdf"
    simple_pdf_file = TXT_DIR / f"{file_id}_simple.pdf"

    for file_path in [
        formatted_docx_file,
        simple_docx_file,
        formatted_pdf_file,
        simple_pdf_file,
    ]:
        if not file_path.exists():
            raise HTTPException(
                status_code=404, detail="Một hoặc nhiều file không tồn tại."
            )

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
        zipf.write(formatted_docx_file, arcname=f"{file_id}_formatted.docx")
        zipf.write(simple_docx_file, arcname=f"{file_id}_simple.docx")
        zipf.write(formatted_pdf_file, arcname=f"{file_id}_formatted.pdf")
        zipf.write(simple_pdf_file, arcname=f"{file_id}_simple.pdf")
    zip_buffer.seek(0)
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename={file_id}_files.zip"},
    )


@router.post("/generate", response_model=FileResponseModel)
async def generate_questions(
    file: UploadFile = File(...),
    exam_subject: str = Form(...),
    exam_duration: str = Form(...),
    num_questions: int = Form(..., ge=1),
    level_1: int = Form(..., ge=0),
    level_2: int = Form(..., ge=0),
    level_3: int = Form(..., ge=0),
    level_4: int = Form(..., ge=0),
    level_5: int = Form(..., ge=0),
    level_6: int = Form(..., ge=0),
    api_key: str = get_api_key,
    db: pymysql.connections.Connection = Depends(get_db),
    current_user_id: int = Depends(get_current_user_id),
):
    """Tạo câu hỏi và câu trả lời từ file tải lên dựa trên yêu cầu cấp độ Bloom."""
    try:
        # Kiểm tra loại file
        file_extension = file.filename.split(".")[-1].lower()
        if file_extension not in ["docx", "pdf", "txt"]:
            raise HTTPException(
                status_code=400,
                detail="Loại file không hợp lệ: chỉ hỗ trợ .docx, .pdf hoặc .txt",
            )

        # Tạo request object
        request = QuestionRequest(
            num_questions=num_questions,
            level_1=level_1,
            level_2=level_2,
            level_3=level_3,
            level_4=level_4,
            level_5=level_5,
            level_6=level_6,
        )

        # Kiểm tra tính hợp lệ của request
        validate_request(request)

        # Lưu file tải lên
        file_id = str(uuid4())
        file_path = save_uploaded_file(file, file_id)

        # Trích xuất văn bản từ file
        extracted_text = extract_text_from_file(file_path, file_extension)

        # Lưu văn bản đã trích xuất (nếu cần)
        text_file_path = TXT_DIR / f"{file_id}_extracted.txt"
        with open(text_file_path, "w", encoding="utf-8") as f:
            f.write(extracted_text)

        # Chia nhỏ nội dung thành các đoạn
        ingestion_instance = Ingestion(embedding_model_name="openai")
        docs = ingestion_instance.process_txt(str(text_file_path), chunk_size=2000)
        segments = [(f"Đoạn {i+1}", doc.page_content) for i, doc in enumerate(docs)]

        if not segments:
            raise HTTPException(
                status_code=400, detail="File không có nội dung để xử lý."
            )

        # Kiểm tra số cấp độ khác 0
        level_counts = [level_1, level_2, level_3, level_4, level_5, level_6]
        num_distinct_levels = sum(1 for count in level_counts if count > 0)

        if len(segments) < num_distinct_levels:
            raise HTTPException(
                status_code=400,
                detail=f"Số đoạn văn ({len(segments)}) không đủ để tạo {num_distinct_levels} cấp độ khác nhau.",
            )

            # Trừ 50 token từ người dùng
        cursor = db.cursor()
        cursor.execute(
            "SELECT balance FROM users WHERE idUser = %s", (current_user_id,)
        )
        row = cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Không tìm thấy người dùng")

        current_balance = row[0] or 0
        cost = 10

        if current_balance < cost:
            raise HTTPException(status_code=400, detail="Không đủ token để tạo đề")

        new_balance = current_balance - cost

        cursor.execute(
            "UPDATE users SET balance = %s WHERE idUser = %s",
            (new_balance, current_user_id),
        )

        # Ghi log giao dịch trừ token
        cursor.execute(
            "INSERT INTO transactions (idUser, change_amount, new_balance) VALUES (%s, %s, %s)",
            (current_user_id, -cost, new_balance),
        )

        db.commit()

        # Tạo câu hỏi và câu trả lời
        qa_result = generate_qa_content(segments, request, extracted_text)

        # Tạo file DOCX với định dạng đầy đủ
        formatted_docx_path = TXT_DIR / f"{file_id}_formatted.docx"
        create_formatted_docx_file(
            qa_result, formatted_docx_path, exam_subject, exam_duration
        )

        # Tạo file DOCX đơn giản (lấy header giống file formatted)
        simple_docx_path = TXT_DIR / f"{file_id}_simple.docx"
        create_simple_docx_file(
            qa_result, simple_docx_path, exam_subject, exam_duration
        )

        # Chuyển đổi 2 file DOCX vừa tạo sang PDF
        formatted_pdf_path = TXT_DIR / f"{file_id}_formatted.pdf"
        simple_pdf_path = TXT_DIR / f"{file_id}_simple.pdf"
        convert_docx_to_pdf(formatted_docx_path, formatted_pdf_path)
        convert_docx_to_pdf(simple_docx_path, simple_pdf_path)

        try:
            cursor = db.cursor()
            cursor.execute(
                "INSERT INTO question_history (idUser, num_questions) VALUES (%s, %s)",
                (current_user_id, num_questions),
            )
            db.commit()
        except Exception as e:
            logger.error(f"Lỗi khi ghi lịch sử câu hỏi: {e}")

        # Trả về FileResponse với 2 đường dẫn file
        return FileResponseModel(
            file_id=file_id,
            original_filename=file.filename,
            num_segments=len(segments),
            formatted_docx_download_url=f"/questions/download/formatted/{file_id}",
            simple_docx_download_url=f"/questions/download/simple/{file_id}",
            questions_and_answers=qa_result.qa_results,
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Lỗi không xác định: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Lỗi không xác định: {str(e)}")
